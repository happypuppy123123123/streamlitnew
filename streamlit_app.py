import streamlit as st
from streamlit_ckeditor import st_ckeditor
import docx
from docx.shared import Inches
import base64
from io import BytesIO

def convert_image_to_base64(image):
    image_stream = BytesIO()
    image.save(image_stream, format='PNG')
    image_stream.seek(0)
    img_base64 = base64.b64encode(image_stream.read()).decode('utf-8')
    return f'<img src="data:image/png;base64,{img_base64}" />'

def docx_to_html(path):
    doc = docx.Document(path)
    html = []

    for para in doc.paragraphs:
        html.append(f'<p>{para.text}</p>')

    for table in doc.tables:
        table_html = '<table>'
        for row in table.rows:
            row_html = '<tr>'
            for cell in row.cells:
                row_html += f'<td>{cell.text}</td>'
            row_html += '</tr>'
            table_html += row_html
        table_html += '</table>'
        html.append(table_html)

    for image in doc.inline_shapes:
        img_base64 = convert_image_to_base64(image.image)
        html.append(f'<img src="data:image/png;base64,{img_base64}" />')

    return ''.join(html)

# Example usage in Streamlit
docx_file = st.file_uploader("Upload a DOCX file", type=["docx"])
if docx_file is not None:
    html_content = docx_to_html(docx_file)

    editor_content = st_ckeditor(value=html_content, key="ckeditor")

    if st.button('Save'):
        st.text("Content saved")
else:
    st.write("Please upload a DOCX file to proceed.")
